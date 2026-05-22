#!/usr/bin/env python3
"""
DocSmith — Generate .docx from Markdown with native OMML equations, Chinese fonts,
table support, equation numbering, cross-references, and multi-column layouts.

Supports:
  - $$...$$ display equations
  - ``` code blocks as equation blocks (line-by-line parsing via latex_utils)
  - Multi-column layouts (e.g., IEEE two-column)
  - Single-column span for title/abstract areas
  - Extended heading numbering: "1.", "I.", "A.", "一、", "(1)", etc.
  - Ordered and unordered lists
  - Image placeholders
  - Abstract and acknowledgment sections
  - Table of Contents (TOC) field

Usage:
    python3 -m scripts.generate_docx --output out.docx --config config.json --content content.md
    (run from DocSmith/ directory)
"""

import json
import re
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .fonts import to_pt, set_run_font, set_para_fmt, add_run
from .omml import add_display_eq
from .parser import parse_markdown, DocContext
from .elements import (setup_page, add_heading, add_table, add_table_caption,
                       process_inline_formatting, _add_simple_text,
                       add_reference_section, set_para_single_column,
                       get_column_width_cm, add_section_body_para, add_toc)
from .latex_utils import unicode_to_latex


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_document(config, md_text):
    doc = Document()
    page_cfg = config.get("page", {})
    setup_page(doc, page_cfg)
    headings_cfg = config.get("headings", {})
    body_cfg = config.get("body", {})
    eq_cfg = config.get("equations", {})
    table_cfg = config.get("table", {})
    ctx = DocContext()
    blocks = parse_markdown(md_text)
    body_font_size_pt = to_pt(body_cfg.get("size", 12))
    column_width_cm = get_column_width_cm(page_cfg)
    is_multi_col = page_cfg.get("columns", 1) > 1

    # ---- Pre-scan: register table and equation labels ----
    tnum = 0
    for btype, level, text, meta in blocks:
        if btype == 'table' and meta.get('label'):
            tnum += 1
            ctx.register_label(meta['label'], tnum, f"tab{tnum}")
        # Register equation labels from [公式:label] captions
        if btype in ('display_math', 'code_eq'):
            if meta and meta.get('label'):
                ctx.equation_counter += 1
                ctx.register_label(meta['label'], ctx.equation_counter,
                                   f"eq{ctx.equation_counter}")
            elif btype == 'code_eq':
                # Count code_eq items to pre-register equation numbers
                items = meta.get('items', [])
                for etype, content, eq_num in items:
                    if etype == 'eq':
                        ctx.equation_counter += 1
                        # Check for \label{eq:...} inside the LaTeX
                        label_match = re.search(r'\\label\{eq:(\S+?)\}', content)
                        if label_match:
                            ctx.register_label(label_match.group(1),
                                               ctx.equation_counter,
                                               f"eq{ctx.equation_counter}")
        elif btype == 'display_math':
            # Check for \label{eq:...} inside LaTeX
            label_match = re.search(r'\\label\{eq:(\S+?)\}', text)
            if label_match:
                # equation_counter will be incremented when rendered
                pass  # handled during rendering

    # Reset equation counter for actual rendering
    ctx.equation_counter = 0

    # ---- TOC (optional) ----
    if config.get("toc", False):
        add_toc(doc)

    # ---- State tracking for multi-column spans ----
    prev_blank = True
    seen_first_heading = False
    current_section = None  # 'abstract', 'acknowledgment', or None
    # Reset list counters
    ctx.list_counters = [0, 0, 0, 0]

    for btype, level, text, meta in blocks:
        if btype == 'blank':
            prev_blank = True
            continue

        # Check meta for section context
        if meta and meta.get('section'):
            if current_section != meta['section']:
                current_section = meta['section']
                # Reset list counters on section change
                ctx.list_counters = [0, 0, 0, 0]

        if btype == 'heading':
            if not prev_blank:
                doc.add_paragraph(style='Normal')
            para = add_heading(doc, level, text, headings_cfg, ctx)

            # First heading = title → span all columns
            if is_multi_col and not seen_first_heading:
                set_para_single_column(para)
                seen_first_heading = True

            # Abstract/acknowledgment headings span all columns
            section_type = meta.get('section') if meta else None
            if is_multi_col and section_type in ('abstract', 'acknowledgment'):
                set_para_single_column(para)
                current_section = section_type
            elif section_type is None:
                current_section = None

            prev_blank = False

        elif btype == 'display_math':
            ctx.equation_counter += 1
            para = doc.add_paragraph(style='Normal')
            eq_label = meta.get('label') if meta else None
            if eq_label:
                ctx.register_label(eq_label, ctx.equation_counter,
                                   f"eq{ctx.equation_counter}")
            # Also check for \label inside the LaTeX
            label_match = re.search(r'\\label\{eq:(\S+?)\}', text)
            if label_match:
                ctx.register_label(label_match.group(1), ctx.equation_counter,
                                   f"eq{ctx.equation_counter}")
            add_display_eq(para, text.strip(), eq_cfg, ctx.equation_counter, ctx,
                           column_width_cm=column_width_cm)
            prev_blank = False

        elif btype == 'code_eq':
            items = meta.get('items', [])
            for etype, content, eq_num in items:
                if etype == 'eq':
                    ctx.equation_counter += 1
                    para = doc.add_paragraph(style='Normal')
                    # Check for \label inside LaTeX
                    label_match = re.search(r'\\label\{eq:(\S+?)\}', content)
                    if label_match:
                        ctx.register_label(label_match.group(1),
                                           ctx.equation_counter,
                                           f"eq{ctx.equation_counter}")
                    display_num = eq_num if eq_num and eq_num.isdigit() else ctx.equation_counter
                    add_display_eq(para, content, eq_cfg,
                                   int(display_num) if isinstance(display_num, str) and display_num.isdigit() else display_num,
                                   ctx, column_width_cm=column_width_cm)
                elif etype == 'text_eq':
                    para = doc.add_paragraph(style='Normal')
                    _add_text_equation(para, content, eq_num, eq_cfg, body_cfg,
                                       column_width_cm=column_width_cm)
            prev_blank = False

        elif btype == 'table':
            lines = meta['lines']
            if len(lines) < 2:
                continue
            def parse_row(r):
                return [c.strip() for c in r.strip('|').split('|')]
            headers = parse_row(lines[0])
            if re.match(r'^[\|\s\-:]+$', lines[1]):
                rows = [parse_row(r) for r in lines[2:]]
            else:
                rows = [parse_row(r) for r in lines[1:]]
            add_table(doc, headers, rows, table_cfg, meta.get('caption', ''),
                       ctx, label=meta.get('label'))
            prev_blank = False

        elif btype == 'image':
            # Image placeholder
            para = doc.add_paragraph(style='Normal')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = para.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            bfont = body_cfg.get("font", "SimSun")
            bfont_w = body_cfg.get("font_west", "Times New Roman")
            bsize = body_cfg.get("size", 12)
            path = meta.get('path', '') if meta else ''
            alt = text if text else 'Image'
            placeholder = f"[Figure: {alt}]"
            if path:
                placeholder += f" ({path})"
            run = para.add_run(placeholder)
            set_run_font(run, bfont, bfont_w, to_pt(bsize), italic=True, color=(0x99, 0x99, 0x99))
            prev_blank = False

        elif btype == 'list_item':
            para = doc.add_paragraph(style='Normal')
            set_para_fmt(para, body_cfg.get("line_spacing", 1.5),
                        first_indent=None, alignment=body_cfg.get("alignment", "justify"))
            bfont = body_cfg.get("font", "SimSun")
            bfont_w = body_cfg.get("font_west", "Times New Roman")
            bsize = body_cfg.get("size", 12)
            prefix = f"{'  ' * level}• "
            run = para.add_run(prefix)
            set_run_font(run, bfont, bfont_w, to_pt(bsize))
            pf = para.paragraph_format
            pf.left_indent = Cm(1.0 + level * 0.5)
            process_inline_formatting(para, text, body_cfg, ctx)

            # Multi-column span for section content
            if is_multi_col and current_section in ('abstract', 'acknowledgment'):
                set_para_single_column(para)

            prev_blank = False

        elif btype == 'ordered_list_item':
            style = meta.get('style', '1.') if meta else '1.'
            para = doc.add_paragraph(style='Normal')
            set_para_fmt(para, body_cfg.get("line_spacing", 1.5),
                        first_indent=None, alignment=body_cfg.get("alignment", "justify"))
            bfont = body_cfg.get("font", "SimSun")
            bfont_w = body_cfg.get("font_west", "Times New Roman")
            bsize = body_cfg.get("size", 12)

            # Reset deeper level counters
            ctx.list_counters[level] += 1
            for i in range(level + 1, len(ctx.list_counters)):
                ctx.list_counters[i] = 0
            num = ctx.list_counters[level]

            # Format prefix based on style
            prefix = _format_list_prefix(style, num)
            indent = f"{'  ' * level}"
            run = para.add_run(f"{indent}{prefix}")
            set_run_font(run, bfont, bfont_w, to_pt(bsize))
            pf = para.paragraph_format
            pf.left_indent = Cm(1.0 + level * 0.5)
            process_inline_formatting(para, text, body_cfg, ctx)

            if is_multi_col and current_section in ('abstract', 'acknowledgment'):
                set_para_single_column(para)

            prev_blank = False

        elif btype == 'paragraph':
            section_type = meta.get('section') if meta else None

            if section_type in ('abstract', 'acknowledgment'):
                # Use section body formatting
                add_section_body_para(doc, text, body_cfg, section_type, ctx)
                # Apply single-column span in multi-col layouts
                if is_multi_col:
                    # Find the last paragraph and apply span
                    set_para_single_column(doc.paragraphs[-1])
            else:
                para = doc.add_paragraph(style='Normal')
                set_para_fmt(para, body_cfg.get("line_spacing", 1.5),
                            body_cfg.get("first_line_indent"),
                            body_cfg.get("alignment", "justify"),
                            font_size_pt=body_font_size_pt)
                process_inline_formatting(para, text, body_cfg, ctx)

            prev_blank = False

        elif btype == 'references':
            add_reference_section(doc, meta['items'], body_cfg)
            prev_blank = False

    return doc


def _format_list_prefix(style, num):
    """Format ordered list prefix from style pattern and counter value.
    style is the markdown pattern (e.g. '1.', '1)', 'a.', '(a)') and num is the
    current counter value (1-indexed)."""
    if style.endswith('.'):
        if style[0].isalpha():
            return f"{_list_letter(num, style[0].isupper())}. "
        return f"{num}. "
    elif style.endswith(')'):
        if style[0].isalpha():
            return f"{_list_letter(num, style[0].isupper())}) "
        return f"{num}) "
    elif style.startswith('('):
        inner = style.strip('()')
        if inner.isalpha():
            return f"({_list_letter(num, inner.isupper())}) "
        return f"({num}) "
    return f"{num}. "


def _list_letter(n, upper=True):
    """Convert integer to letter: 1→A/a, 2→B/b, ..., 26→Z/z, 27→AA/aa."""
    result = ''
    while n > 0:
        n, rem = divmod(n - 1, 26)
        result = chr(ord('A' if upper else 'a') + rem) + result
    return result


def _add_text_equation(para, text, eq_num, eq_cfg, body_cfg, column_width_cm=None):
    """Render a text constraint as a centered equation-like line with optional number."""
    if column_width_cm is None:
        column_width_cm = 15.92  # A4 single-column fallback

    bfont = body_cfg.get("font", "SimSun")
    bfont_w = body_cfg.get("font_west", "Times New Roman")
    bsize = body_cfg.get("size", 12)
    nfont = eq_cfg.get("numbering_font", "Times New Roman")
    nsize = to_pt(eq_cfg.get("numbering_size", 12))

    if eq_num:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        center_stop = Cm(column_width_cm / 2.0)
        right_stop = Cm(column_width_cm)
        pf.tab_stops.add_tab_stop(center_stop, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        pf.tab_stops.add_tab_stop(right_stop, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        para.add_run("\t")
        run = para.add_run(text)
        set_run_font(run, bfont, bfont_w, to_pt(bsize), italic=True)
        run2 = para.add_run(f"\t({eq_num})")
        set_run_font(run2, nfont, nfont, nsize)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_run_font(run, bfont, bfont_w, to_pt(bsize), italic=True)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    p = argparse.ArgumentParser(description="DocSmith — MD+LaTeX → professional .docx")
    p.add_argument("--output", required=True)
    p.add_argument("--config", required=True)
    p.add_argument("--content", required=True)
    args = p.parse_args()
    with open(args.config, 'r', encoding='utf-8') as f:
        config = json.load(f)
    with open(args.content, 'r', encoding='utf-8') as f:
        md = f.read()
    doc = build_document(config, md)
    doc.save(args.output)
    print(f"OK — {args.output}")


if __name__ == "__main__":
    main()
