"""Document elements: page setup, headings, tables, inline formatting.

Supports:
  - Single and multi-column page layouts (via OXML w:cols)
  - Single-column paragraph spans (w:cnt) for title/abstract areas
  - Column-width-aware tab stops for equation numbering
  - Auto-numbered headings with configurable numbering styles:
    "1.", "1.1", "一、", "（一）", "I.", "II.", "A.", "B."
  - Tables with captions and cross-reference bookmarks
  - Reference sections with hanging indents
  - Abstract and acknowledgment sections with special formatting
"""

import re
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .fonts import to_pt, set_run_font, set_para_fmt, add_run
from .equations import append_omml, add_internal_hyperlink, add_bookmark_to_para
from .numbering import heading_prefix
from .parser import resolve_cross_ref


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def setup_page(doc, cfg):
    """Configure page size, margins, and column layout for all sections.

    cfg keys:
      size        : "A4" or "letter" (default A4)
      margin_*    : top, bottom, left, right in cm (default 2.54)
      columns     : None (single column, default) or int (2 for two-column)
      column_gap  : gap between columns in cm (default 0.63 for 0.25 inches)

    Two-column layout uses OXML w:cols injection since python-docx
    does not natively support multi-column sections.
    """
    for s in doc.sections:
        smap = {"A4": (Cm(21.0), Cm(29.7)), "letter": (Cm(21.59), Cm(27.94))}
        if cfg.get("size") in smap:
            s.page_width, s.page_height = smap[cfg["size"]]
        s.top_margin = Cm(cfg.get("margin_top", 2.54))
        s.bottom_margin = Cm(cfg.get("margin_bottom", 2.54))
        s.left_margin = Cm(cfg.get("margin_left", 2.54))
        s.right_margin = Cm(cfg.get("margin_right", 2.54))

        # Multi-column layout
        num_cols = cfg.get("columns", None)
        if num_cols and num_cols > 1:
            sectPr = s._sectPr
            cols = OxmlElement('w:cols')
            cols.set(qn('w:num'), str(num_cols))
            cols.set(qn('w:equalWidth'), '1')
            # Column gap: default 0.63cm (~0.25 inch)
            gap_cm = cfg.get("column_gap", 0.63)
            gap_twips = str(int(gap_cm * 567))  # 1 cm ≈ 567 twips
            cols.set(qn('w:space'), gap_twips)
            for e in sectPr.findall(qn('w:cols')):
                sectPr.remove(e)
            sectPr.append(cols)


def get_column_width_cm(cfg):
    """Compute usable column width in cm from page config.
    Used to calculate correct equation tab stops."""
    page_size = cfg.get("size", "A4")
    ml = cfg.get("margin_left", 2.54)
    mr = cfg.get("margin_right", 2.54)
    num_cols = cfg.get("columns", 1) or 1
    col_gap = cfg.get("column_gap", 0.63)

    page_widths = {"A4": 21.0, "letter": 21.59}
    full_width = page_widths.get(page_size, 21.0)
    usable = full_width - ml - mr

    if num_cols > 1:
        return (usable - col_gap * (num_cols - 1)) / num_cols
    return usable


def set_para_single_column(para):
    """Force paragraph to span all columns (for title, authors, abstract).

    Uses the w:cnt (continuous) element to break out of multi-column flow.
    Paragraphs after this one will resume normal column flow.
    """
    pPr = para._p.get_or_add_pPr()
    # Remove any existing cnt element to avoid duplicates
    for old in pPr.findall(qn('w:cnt')):
        pPr.remove(old)
    cnt = OxmlElement('w:cnt')
    pPr.append(cnt)


# ---------------------------------------------------------------------------
# Heading
# ---------------------------------------------------------------------------

def add_heading(doc, level, text, headings_cfg, ctx):
    """Add a heading paragraph bound to Word's Heading N style.
    Forces black font color (overrides Word's default blue) and no indent."""
    key = f"h{level}"
    hc = headings_cfg.get(key)
    if not hc:
        return
    prefix = heading_prefix(level, headings_cfg, ctx)
    style_name = f"Heading {level}"
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(prefix + text)
    set_run_font(run, hc.get("font", "SimHei"), hc.get("font_west", "Arial"),
                 to_pt(hc.get("size", 16)), bold=hc.get("bold", True),
                 color=(0, 0, 0))
    pf = para.paragraph_format
    pf.line_spacing = hc.get("line_spacing", 1.5)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    return para


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def add_table(doc, headers, rows, table_cfg, caption_text, ctx, label=None):
    """Create a formatted Word table with optional caption."""
    ctx.table_counter += 1
    tnum = ctx.table_counter
    anchor = f"tab{tnum}"

    if label:
        ctx.register_label(label, tnum, anchor)

    # Caption (above table)
    if caption_text and table_cfg.get("caption_position", "above") == "above":
        add_table_caption(doc, caption_text, table_cfg, tnum, ctx, anchor)

    # Build table
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'

    # Header row
    hfont = table_cfg.get("header_font", "SimHei")
    hfont_w = table_cfg.get("header_font_west", "Arial")
    hsize = to_pt(table_cfg.get("header_size", "小四"))
    hbold = table_cfg.get("header_bold", True)
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        set_run_font(run, hfont, hfont_w, hsize, bold=hbold)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    bfont = table_cfg.get("body_font", "SimSun")
    bfont_w = table_cfg.get("body_font_west", "Times New Roman")
    bsize = to_pt(table_cfg.get("body_size", "小五"))

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            if j >= ncols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            if '$' in str(val):
                _add_simple_text(p, str(val), bfont, bfont_w, bsize)
            else:
                run = p.add_run(str(val))
                set_run_font(run, bfont, bfont_w, bsize)

    # Caption (below table)
    if caption_text and table_cfg.get("caption_position", "above") != "above":
        add_table_caption(doc, caption_text, table_cfg, tnum, ctx, anchor)

    return table


def add_table_caption(doc, caption_text, table_cfg, tnum, ctx, anchor=None):
    """Add '表 1 标题文本' caption paragraph."""
    if anchor is None:
        anchor = f"tab{tnum}"
    num_fmt = table_cfg.get("numbering", "表 {n} ")
    prefix = num_fmt.replace("{n}", str(tnum))
    para = doc.add_paragraph(style='Caption')
    run = para.add_run(prefix + caption_text)
    cfont = table_cfg.get("caption_font", "黑体")
    cfont_w = table_cfg.get("caption_font_west", "Arial")
    csize = to_pt(table_cfg.get("caption_size", "小五"))
    cbold = table_cfg.get("caption_bold", True)
    set_run_font(run, cfont, cfont_w, csize, bold=cbold, color=(0, 0, 0))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bookmark_to_para(para, anchor, ctx.next_bookmark(anchor))
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Inline formatting (cross-refs, LaTeX, bold/italic)
# ---------------------------------------------------------------------------

def _add_simple_text(para, text, font_cn, font_west, size, bold=False, italic=False):
    """Add text with optional inline LaTeX ($...$). size may be a float (pt) or
    a string like '小五' — to_pt handles both idempotently."""
    parts = re.split(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', text)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            if part:
                run = para.add_run(part)
                set_run_font(run, font_cn, font_west, to_pt(size), bold, italic)
        else:
            if part.strip():
                append_omml(para, part.strip(), display=False)


def process_inline_formatting(para, text, body_cfg, ctx):
    """Process cross-refs, inline LaTeX, citations, and bold/italic.
    Delegates to InlineProcessor pipeline (lazy import to avoid circular deps)."""
    from .renderer import InlineProcessor
    InlineProcessor.process(para, text, body_cfg, ctx)


# ---------------------------------------------------------------------------
# Abstract / Acknowledgment sections
# ---------------------------------------------------------------------------

def add_section_body_para(doc, text, body_cfg, section_type, ctx):
    """Add a body paragraph within an abstract or acknowledgment section.
    Uses body text formatting. The section heading already provides the label,
    so no prefix is prepended here."""
    para = doc.add_paragraph(style='Normal')
    bfont = body_cfg.get("font", "SimSun")
    bfont_w = body_cfg.get("font_west", "Times New Roman")
    bsize = body_cfg.get("size", 12)
    body_font_size_pt = to_pt(bsize)

    set_para_fmt(para, body_cfg.get("line_spacing", 1.5),
                 body_cfg.get("first_line_indent"),
                 body_cfg.get("alignment", "justify"),
                 font_size_pt=body_font_size_pt)

    process_inline_formatting(para, text, body_cfg, ctx)
    return para


# ---------------------------------------------------------------------------
# References (GB/T 7714)
# ---------------------------------------------------------------------------

def add_reference_section(doc, ref_items, body_cfg, headings_cfg=None, ref_heading_level=2):
    """Add a reference list. Heading level is configurable (default H2).
    Each reference is a Normal paragraph with hanging indent."""
    if headings_cfg is None:
        headings_cfg = {
            "h2": {"font": "SimHei", "font_west": "Arial", "size": "三号", "bold": True}
        }
    key = f"h{ref_heading_level}"
    hc = headings_cfg.get(key, {"font": "SimHei", "font_west": "Arial", "size": "三号", "bold": True})

    from .context import DocContext
    ctx = DocContext()
    style_name = f"Heading {ref_heading_level}"
    para = doc.add_paragraph(style=style_name)
    run = para.add_run("参考文献")
    set_run_font(run, hc.get("font", "SimHei"), hc.get("font_west", "Arial"),
                 to_pt(hc.get("size", 16)), bold=hc.get("bold", True),
                 color=(0, 0, 0))
    pf = para.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)

    # Each reference as Normal paragraph with hanging indent
    for item in ref_items:
        ref_para = doc.add_paragraph(style='Normal')
        bfont = body_cfg.get("font", "SimSun")
        bfont_w = body_cfg.get("font_west", "Times New Roman")
        bsize_val = body_cfg.get("size", "小五")
        bsize = to_pt(bsize_val)
        run = ref_para.add_run(item)
        set_run_font(run, bfont, bfont_w, bsize, color=(0, 0, 0))
        rpf = ref_para.paragraph_format
        rpf.line_spacing = 1.5
        rpf.space_before = Pt(0)
        rpf.space_after = Pt(0)
        rpf.first_line_indent = Pt(0)
        # Hanging indent: left indent = 2 chars, first line = -2 chars
        hang = Pt(bsize * 2)
        rpf.left_indent = hang
        rpf.first_line_indent = -hang


# ---------------------------------------------------------------------------
# TOC (Table of Contents)
# ---------------------------------------------------------------------------

def add_toc(doc):
    """Insert a Word TOC field at the current cursor position.
    The TOC will populate when the user opens the document in Word
    and confirms 'Update Table of Contents'."""
    para = doc.add_paragraph(style='Normal')
    run = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z '
    run2._r.append(instrText)

    run3 = para.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar_separate)

    run4 = para.add_run('[ Table of Contents — right-click and select "Update Field" in Word ]')
    set_run_font(run4, "SimSun", "Times New Roman", to_pt(10), italic=True)

    run5 = para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar_end)

    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    return para
