"""OMML equation rendering and bookmark/hyperlink utilities."""

from lxml import etree
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .fonts import to_pt, set_run_font

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

try:
    from latex2word import LatexToWordElement
    HAS_LATEX2WORD = True
except ImportError:
    HAS_LATEX2WORD = False
    import sys
    print("DocSmith: latex2word not installed. Equations will show as red placeholders.",
          file=sys.stderr)
    print("  Install with: pip3 install latex2word", file=sys.stderr)


def build_ommath_para(omml_element):
    """Wrap an m:oMath element in a properly structured m:oMathPara (display eq)."""
    omp = etree.Element(f"{{{MATH_NS}}}oMathPara")
    omp_pr = etree.SubElement(omp, f"{{{MATH_NS}}}oMathParaPr")
    jc = etree.SubElement(omp_pr, f"{{{MATH_NS}}}jc")
    jc.set(f"{{{MATH_NS}}}val", "center")
    om = etree.SubElement(omp, f"{{{MATH_NS}}}oMath")
    for child in omml_element:
        om.append(child)
    return omp


def append_omml(para, latex_str, display=True):
    """Convert LaTeX to OMML and append to paragraph.
    Returns the OMML element (oMathPara for display, oMath for inline)."""
    if not HAS_LATEX2WORD:
        run = para.add_run(f"[Equation: {latex_str}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return None
    try:
        eq = LatexToWordElement(latex_str)
        omml = eq.element()
        if display:
            omp = build_ommath_para(omml)
            para._element.append(omp)
            return omp
        else:
            para._element.append(omml)
            return omml
    except Exception:
        run = para.add_run(f"[Eq: {latex_str[:60]}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return None


def _compute_tab_stops(column_width_cm):
    """Compute center and right tab stops for equation numbering layout.
    column_width_cm is the usable width of a single column (after margins)."""
    center_tab = column_width_cm / 2.0
    right_tab = column_width_cm
    return Cm(center_tab), Cm(right_tab)


def add_display_eq(para, latex_str, eq_cfg, eq_num, ctx, column_width_cm=None):
    """Add a display equation.
    Numbered: tab-stop layout with equation centered, number right-aligned.
    Unnumbered: paragraph centered, OMML only.

    column_width_cm is the usable column width (page width minus margins,
    divided by number of columns minus gaps). If None, falls back to A4
    single-column default (15.92cm).
    """
    if column_width_cm is None:
        column_width_cm = 15.92  # A4 single-column fallback

    if eq_cfg.get("numbering"):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        center_stop, right_stop = _compute_tab_stops(column_width_cm)
        pf.tab_stops.add_tab_stop(center_stop, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        pf.tab_stops.add_tab_stop(right_stop, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        para.add_run("\t")
        append_omml(para, latex_str.strip(), display=True)
        fmt = eq_cfg.get("numbering_format", "({n})")
        num_text = fmt.replace("{n}", str(eq_num))
        run = para.add_run(f"\t{num_text}")
        nfont = eq_cfg.get("numbering_font", "Times New Roman")
        nsize = to_pt(eq_cfg.get("numbering_size", 12))
        set_run_font(run, nfont, nfont, nsize)
        bm_id = ctx.next_bookmark(f"eq{eq_num}")
        add_bookmark_to_run(run, f"eq{eq_num}", bm_id)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_omml(para, latex_str.strip(), display=True)


def add_bookmark_to_run(run, name, bmid):
    """Wrap a run element with bookmark start/end tags."""
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bmid))
    bs.set(qn('w:name'), name)
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bmid))
    run._r.addprevious(bs)
    run._r.addnext(be)


def add_bookmark_to_para(para, name, bmid):
    """Add bookmark start/end to a paragraph element."""
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bmid))
    bs.set(qn('w:name'), name)
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bmid))
    para._p.append(bs)
    para._p.append(be)


def add_internal_hyperlink(para, anchor, display_text, font_cn, font_west, size):
    """Add a clickable internal cross-reference (e.g. '表 1')."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rs = OxmlElement('w:rStyle')
    rs.set(qn('w:val'), 'Hyperlink')
    rPr.append(rs)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(round(to_pt(size) * 2)))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = display_text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    hl.append(r)
    para._p.append(hl)
