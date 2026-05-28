"""Block renderer: dispatches parsed blocks to docx element functions.

Extension points:
  - BlockRenderer.register_block_type — decorator for adding new block types
  - InlineProcessor.register — decorator for adding inline content processors
  - _add_processed_run — renders processed inline segments to runs
"""

import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .fonts import to_pt, set_run_font, set_para_fmt
from .equations import add_display_eq, append_omml, add_internal_hyperlink, add_bookmark_to_para
from .elements import (
    add_heading, add_table, add_reference_section, set_para_single_column,
    add_section_body_para
)
from .numbering import format_list_prefix
from .parser import resolve_cross_ref


# ===========================================================================
# Block Renderer Registry
# ===========================================================================

# ---- Block renderer registry (module-level for decorator use at class definition time) ----

_block_renderers = {}


def register_block_type(btype):
    """Decorator: register a renderer method for a block type."""
    def decorator(func):
        _block_renderers[btype] = func
        return func
    return decorator


class BlockRenderer:
    """Renders parsed blocks into a python-docx Document.

    Uses a class-level dispatch dict. New block types register via the
    @register_block_type('type_name') decorator.
    """

    _block_renderers = _block_renderers

    def __init__(self, doc, config, ctx):
        self.doc = doc
        self.config = config
        self.ctx = ctx
        self.body_cfg = config.get("body", {})
        self.eq_cfg = config.get("equations", {})
        self.table_cfg = config.get("table", {})
        self.headings_cfg = config.get("headings", {})
        self.page_cfg = config.get("page", {})

        self.body_font_size_pt = to_pt(self.body_cfg.get("size", 12))
        self.column_width_cm = _get_column_width_cm(self.page_cfg)
        self.is_multi_col = self.page_cfg.get("columns", 1) > 1

        self.prev_blank = True
        self.seen_first_heading = False
        self.current_section = None

    def render(self, block):
        """Dispatch a single block to its registered renderer."""
        btype = block[0]
        renderer = self._block_renderers.get(btype)
        if renderer:
            renderer(self, block)

    # ---- Built-in block renderers ----

    @register_block_type('blank')
    def _render_blank(self, block):
        self.prev_blank = True

    @register_block_type('heading')
    def _render_heading(self, block):
        _, level, text, meta = block
        if not self.prev_blank:
            self.doc.add_paragraph(style='Normal')
        para = add_heading(self.doc, level, text, self.headings_cfg, self.ctx)

        if self.is_multi_col and not self.seen_first_heading:
            set_para_single_column(para)
            self.seen_first_heading = True

        section_type = meta.get('section') if meta else None
        if self.is_multi_col and section_type in ('abstract', 'acknowledgment'):
            set_para_single_column(para)
            self.current_section = section_type
        elif section_type is None:
            self.current_section = None

        self.prev_blank = False

    @register_block_type('display_math')
    def _render_display_math(self, block):
        _, _, text, meta = block
        self.ctx.equation_counter += 1
        para = self.doc.add_paragraph(style='Normal')
        eq_label = meta.get('label') if meta else None
        if eq_label:
            self.ctx.register_label(eq_label, self.ctx.equation_counter,
                                    f"eq{self.ctx.equation_counter}")
        label_match = re.search(r'\\label\{eq:(\S+?)\}', text)
        if label_match:
            self.ctx.register_label(label_match.group(1), self.ctx.equation_counter,
                                    f"eq{self.ctx.equation_counter}")
        add_display_eq(para, text.strip(), self.eq_cfg, self.ctx.equation_counter,
                       self.ctx, column_width_cm=self.column_width_cm)
        self.prev_blank = False

    @register_block_type('code_eq')
    def _render_code_eq(self, block):
        _, _, _, meta = block
        items = meta.get('items', [])
        for etype, content, eq_num in items:
            if etype == 'eq':
                self.ctx.equation_counter += 1
                para = self.doc.add_paragraph(style='Normal')
                label_match = re.search(r'\\label\{eq:(\S+?)\}', content)
                if label_match:
                    self.ctx.register_label(label_match.group(1),
                                            self.ctx.equation_counter,
                                            f"eq{self.ctx.equation_counter}")
                display_num = eq_num if eq_num and eq_num.isdigit() else self.ctx.equation_counter
                add_display_eq(para, content, self.eq_cfg,
                               int(display_num) if isinstance(display_num, str) and display_num.isdigit() else display_num,
                               self.ctx, column_width_cm=self.column_width_cm)
            elif etype == 'text_eq':
                para = self.doc.add_paragraph(style='Normal')
                _add_text_equation(para, content, eq_num, self.eq_cfg, self.body_cfg,
                                   column_width_cm=self.column_width_cm)
        self.prev_blank = False

    @register_block_type('table')
    def _render_table(self, block):
        _, _, _, meta = block
        lines = meta['lines']
        if len(lines) < 2:
            return

        def parse_row(r):
            return [c.strip() for c in r.strip('|').split('|')]
        headers = parse_row(lines[0])
        if re.match(r'^[\|\s\-:]+$', lines[1]):
            rows = [parse_row(r) for r in lines[2:]]
        else:
            rows = [parse_row(r) for r in lines[1:]]
        add_table(self.doc, headers, rows, self.table_cfg, meta.get('caption', ''),
                  self.ctx, label=meta.get('label'))
        self.prev_blank = False

    @register_block_type('table_caption')
    def _render_table_caption(self, block):
        self.prev_blank = False

    @register_block_type('image')
    def _render_image(self, block):
        _, _, text, meta = block
        para = self.doc.add_paragraph(style='Normal')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        bfont = self.body_cfg.get("font", "SimSun")
        bfont_w = self.body_cfg.get("font_west", "Times New Roman")
        bsize = self.body_cfg.get("size", 12)
        path = meta.get('path', '') if meta else ''
        alt = text if text else 'Image'
        placeholder = f"[Figure: {alt}]"
        if path:
            placeholder += f" ({path})"
        run = para.add_run(placeholder)
        set_run_font(run, bfont, bfont_w, to_pt(bsize), italic=True, color=(0x99, 0x99, 0x99))
        self.prev_blank = False

    @register_block_type('figure_caption')
    def _render_figure_caption(self, block):
        """Render a figure caption with auto-numbering and cross-reference bookmark."""
        _, _, _, meta = block
        self.ctx.figure_counter += 1
        fnum = self.ctx.figure_counter
        label = meta.get('label')
        if label:
            self.ctx.register_label(label, fnum, f"fig{fnum}")
        caption = meta.get('caption', '')
        para = self.doc.add_paragraph(style='Normal')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bfont = self.body_cfg.get("font", "SimSun")
        bfont_w = self.body_cfg.get("font_west", "Times New Roman")
        bsize = self.body_cfg.get("size", 12)
        prefix = f"图 {fnum}"
        if caption:
            prefix += f"  {caption}"
        run = para.add_run(prefix)
        set_run_font(run, bfont, bfont_w, to_pt(bsize), bold=True)
        add_bookmark_to_para(para, f"fig{fnum}", self.ctx.next_bookmark(f"fig{fnum}"))
        self.prev_blank = False

    @register_block_type('list_item')
    def _render_list_item(self, block):
        _, level, text, meta = block
        para = self.doc.add_paragraph(style='Normal')
        set_para_fmt(para, self.body_cfg.get("line_spacing", 1.5),
                     first_indent=None, alignment=self.body_cfg.get("alignment", "justify"))
        bfont = self.body_cfg.get("font", "SimSun")
        bfont_w = self.body_cfg.get("font_west", "Times New Roman")
        bsize = self.body_cfg.get("size", 12)
        prefix = f"{'  ' * level}• "
        run = para.add_run(prefix)
        set_run_font(run, bfont, bfont_w, to_pt(bsize))
        pf = para.paragraph_format
        pf.left_indent = Cm(1.0 + level * 0.5)
        process_inline_formatting(para, text, self.body_cfg, self.ctx)
        if self.is_multi_col and self.current_section in ('abstract', 'acknowledgment'):
            set_para_single_column(para)
        self.prev_blank = False

    @register_block_type('ordered_list_item')
    def _render_ordered_list_item(self, block):
        _, level, text, meta = block
        style = meta.get('style', '1.') if meta else '1.'
        para = self.doc.add_paragraph(style='Normal')
        set_para_fmt(para, self.body_cfg.get("line_spacing", 1.5),
                     first_indent=None, alignment=self.body_cfg.get("alignment", "justify"))
        bfont = self.body_cfg.get("font", "SimSun")
        bfont_w = self.body_cfg.get("font_west", "Times New Roman")
        bsize = self.body_cfg.get("size", 12)

        self.ctx.list_counters[level] += 1
        for i in range(level + 1, len(self.ctx.list_counters)):
            self.ctx.list_counters[i] = 0
        num = self.ctx.list_counters[level]
        prefix = format_list_prefix(style, num)
        indent = f"{'  ' * level}"
        run = para.add_run(f"{indent}{prefix}")
        set_run_font(run, bfont, bfont_w, to_pt(bsize))
        pf = para.paragraph_format
        pf.left_indent = Cm(1.0 + level * 0.5)
        process_inline_formatting(para, text, self.body_cfg, self.ctx)
        if self.is_multi_col and self.current_section in ('abstract', 'acknowledgment'):
            set_para_single_column(para)
        self.prev_blank = False

    @register_block_type('paragraph')
    def _render_paragraph(self, block):
        _, _, text, meta = block
        section_type = meta.get('section') if meta else None
        if section_type in ('abstract', 'acknowledgment'):
            add_section_body_para(self.doc, text, self.body_cfg, section_type, self.ctx)
            if self.is_multi_col:
                set_para_single_column(self.doc.paragraphs[-1])
        else:
            para = self.doc.add_paragraph(style='Normal')
            set_para_fmt(para, self.body_cfg.get("line_spacing", 1.5),
                         self.body_cfg.get("first_line_indent"),
                         self.body_cfg.get("alignment", "justify"),
                         font_size_pt=self.body_font_size_pt)
            process_inline_formatting(para, text, self.body_cfg, self.ctx)
        self.prev_blank = False

    @register_block_type('references')
    def _render_references(self, block):
        _, _, _, meta = block
        ref_heading_level = self.config.get("references", {}).get("heading_level", 2)
        ref_items = meta.get('items', [])

        # If citation_map has entries, reorder references to match citation order
        if self.ctx.citation_map:
            ordered = [None] * len(self.ctx.citation_map)
            for item in ref_items:
                key_match = re.match(r'^\[@(\w+)\]\s*(.+)', item)
                if key_match:
                    key = key_match.group(1)
                    text = key_match.group(2)
                    if key in self.ctx.citation_map:
                        idx = self.ctx.citation_map[key] - 1
                        if idx < len(ordered):
                            ordered[idx] = text
            ref_items = [r for r in ordered if r is not None] or ref_items
        else:
            # Strip [@key] prefixes from items that have them
            ref_items = [re.sub(r'^\[@\w+\]\s*', '', item) for item in ref_items]

        add_reference_section(self.doc, ref_items, self.body_cfg,
                              headings_cfg=self.headings_cfg,
                              ref_heading_level=ref_heading_level)
        self.prev_blank = False


# ===========================================================================
# Inline Processor Pipeline
# ===========================================================================

class InlineProcessor:
    """Ordered pipeline for processing inline content.

    Processors are registered with a priority (lower = runs first).
    Each processor function receives (text, body_cfg, ctx) and returns either:
      - None if no match
      - (consumed_chars, [run_data_dict, ...]) if matched

    Each run_data_dict can have:
      - {'text': str, 'bold': bool, 'italic': bool, 'superscript': bool}
      - {'omml': str} for inline LaTeX
      - {'hyperlink': str, 'anchor': str, 'display': str}
    """

    _processors = []  # list of (name, priority, func)

    @classmethod
    def register(cls, name, priority=100):
        """Decorator: register an inline processor. Lower priority runs first."""
        def decorator(func):
            cls._processors.append((name, priority, func))
            cls._processors.sort(key=lambda x: x[1])
            return func
        return decorator

    @classmethod
    def process(cls, para, text, body_cfg, ctx):
        """Run all registered processors on text, appending runs to para."""
        remaining = text
        bfont = body_cfg.get("font", "SimSun")
        bfont_w = body_cfg.get("font_west", "Times New Roman")
        bsize = body_cfg.get("size", 12)

        while remaining:
            matched = False
            for name, priority, func in cls._processors:
                result = func(remaining, body_cfg, ctx)
                if result is not None:
                    consumed, runs_data = result
                    for rd in runs_data:
                        _add_processed_run(para, rd, bfont, bfont_w, bsize)
                    remaining = remaining[consumed:]
                    matched = True
                    break
            if not matched:
                # Pass through one char as plain text
                run = para.add_run(remaining[0])
                set_run_font(run, bfont, bfont_w, to_pt(bsize))
                remaining = remaining[1:]


def _add_processed_run(para, data, font_cn, font_west, size):
    """Add a run from processed inline data to the paragraph."""
    if 'omml' in data:
        append_omml(para, data['omml'], display=False)
    elif 'hyperlink' in data:
        add_internal_hyperlink(para, data['anchor'], data['display'],
                               font_cn, font_west, size)
    else:
        text = data.get('text', '')
        bold = data.get('bold', False)
        italic = data.get('italic', False)
        superscript = data.get('superscript', False)
        run = para.add_run(text)
        set_run_font(run, font_cn, font_west, to_pt(size), bold=bold, italic=italic)
        if superscript:
            run.font.superscript = True


# ---- Built-in inline processors ----

@InlineProcessor.register('cross_ref', priority=10)
def _match_cross_ref(text, body_cfg, ctx):
    """Match \\ref{tab:...}, \\ref{eq:...}, or \\ref{fig:...}"""
    m = re.match(r'\\ref\{(tab|eq|fig):(\S+?)\}', text)
    if not m:
        return None
    ref_type, ref_id = m.group(1), m.group(2)
    disp, anchor = resolve_cross_ref(m.group(0), body_cfg, ctx)
    if disp:
        return len(m.group(0)), [{'hyperlink': True, 'anchor': anchor, 'display': disp}]
    return None


@InlineProcessor.register('citation', priority=15)
def _match_citation(text, body_cfg, ctx):
    """Match [@key1, @key2] citation patterns. Renders as superscript numbers."""
    m = re.match(r'\[((?:@\w+(?:,\s*)?)+)\]', text)
    if not m:
        return None
    keys = re.findall(r'@(\w+)', m.group(1))
    nums = []
    for key in keys:
        if key not in ctx.citation_map:
            ctx.citation_map[key] = len(ctx.citation_map) + 1
        nums.append(str(ctx.citation_map[key]))
    display = f"[{','.join(nums)}]"
    return len(m.group(0)), [{'text': display, 'superscript': True}]


@InlineProcessor.register('inline_math', priority=20)
def _match_inline_math(text, body_cfg, ctx):
    """Match $...$ inline LaTeX."""
    m = re.match(r'\$(.+?)\$', text)
    if not m:
        return None
    return len(m.group(0)), [{'omml': m.group(1)}]


@InlineProcessor.register('bold_italic', priority=30)
def _match_formatting(text, body_cfg, ctx):
    """Match ***bold-italic***, **bold**, or *italic*."""
    m = re.match(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*', text)
    if not m:
        return None
    content = m.group(1) or m.group(2) or m.group(3)
    bold = bool(m.group(1) or m.group(2))
    italic = bool(m.group(1) or m.group(3))
    return len(m.group(0)), [{'text': content, 'bold': bold, 'italic': italic}]


# ===========================================================================
# Public API — replaces old process_inline_formatting
# ===========================================================================

def process_inline_formatting(para, text, body_cfg, ctx):
    """Process inline content through the processor pipeline."""
    InlineProcessor.process(para, text, body_cfg, ctx)


# ===========================================================================
# Helpers
# ===========================================================================

def _get_column_width_cm(cfg):
    """Compute usable column width in cm from page config."""
    page_size = cfg.get("size", "A4")
    ml = cfg.get("margin_left", 2.54)
    mr = cfg.get("margin_right", 2.54)
    num_cols = cfg.get("columns", 1) or 1
    col_gap = cfg.get("column_gap", 0.63)
    page_widths = {"A4": 21.0, "letter": 21.59}
    full_width = page_widths.get(page_size, 21.0)
    usable = full_width - ml - mr
    if num_cols > 1:
        return (usable - col_gap * (num_cols - 1)) / num_cols
    return usable


def _add_text_equation(para, text, eq_num, eq_cfg, body_cfg, column_width_cm=None):
    """Render a text constraint as a centered equation-like line with optional number."""
    if column_width_cm is None:
        column_width_cm = 15.92
    bfont = body_cfg.get("font", "SimSun")
    bfont_w = body_cfg.get("font_west", "Times New Roman")
    bsize = body_cfg.get("size", 12)
    nfont = eq_cfg.get("numbering_font", "Times New Roman")
    nsize = to_pt(eq_cfg.get("numbering_size", 12))
    if eq_num:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        center_stop = Cm(column_width_cm / 2.0)
        right_stop = Cm(column_width_cm)
        pf.tab_stops.add_tab_stop(center_stop, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        pf.tab_stops.add_tab_stop(right_stop, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        para.add_run("\t")
        run = para.add_run(text)
        set_run_font(run, bfont, bfont_w, to_pt(bsize), italic=True)
        run2 = para.add_run(f"\t({eq_num})")
        set_run_font(run2, nfont, nfont, nsize)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_run_font(run, bfont, bfont_w, to_pt(bsize), italic=True)
