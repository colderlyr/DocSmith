"""LaTeX equation pipeline: Unicode→LaTeX preprocessing + LaTeX→OMML rendering.

Pipeline:
  raw text → unicode_to_latex() → clean LaTeX → append_omml() / add_display_eq()
"""

import re
import sys
from lxml import etree
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .fonts import to_pt, set_run_font


# ===========================================================================
# latex2word availability
# ===========================================================================

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

try:
    from latex2word import LatexToWordElement
    HAS_LATEX2WORD = True
except ImportError:
    HAS_LATEX2WORD = False
    print("DocSmith: latex2word not installed. Equations will show as red placeholders.",
          file=sys.stderr)
    print("  Install with: pip3 install latex2word", file=sys.stderr)


# ===========================================================================
# Unicode → LaTeX symbol mappings
# ===========================================================================

# Greek and math symbols (sorted longest-first for safe replacement)
UNICODE_MATH_SYMBOLS = [
    # Operators / relations
    ('∥', r'\parallel '), ('…', r'\dots '), ('∑', r'\sum '), ('∏', r'\prod '),
    ('√', r'\sqrt '), ('⟨', r'\langle '), ('⟩', r'\rangle '),
    ('∇', r'\nabla '), ('∂', r'\partial '), ('∫', r'\int '),
    # Greek lowercase
    ('α', r'\alpha '), ('β', r'\beta '), ('γ', r'\gamma '),
    ('δ', r'\delta '), ('ε', r'\epsilon '), ('ζ', r'\zeta '),
    ('η', r'\eta '), ('θ', r'\theta '), ('ι', r'\iota '),
    ('κ', r'\kappa '), ('λ', r'\lambda '), ('μ', r'\mu '),
    ('ν', r'\nu '), ('ξ', r'\xi '), ('π', r'\pi '),
    ('ρ', r'\rho '), ('σ', r'\sigma '), ('τ', r'\tau '),
    ('υ', r'\upsilon '), ('φ', r'\phi '), ('χ', r'\chi '),
    ('ψ', r'\psi '), ('ω', r'\omega '),
    # Greek uppercase
    ('Γ', r'\Gamma '), ('Δ', r'\Delta '), ('Θ', r'\Theta '),
    ('Λ', r'\Lambda '), ('Ξ', r'\Xi '), ('Π', r'\Pi '),
    ('Σ', r'\Sigma '), ('Φ', r'\Phi '), ('Ψ', r'\Psi '),
    ('Ω', r'\Omega '),
    # Operators
    ('·', r'\cdot '), ('×', r'\times '),
    ('≤', r'\leq '), ('≥', r'\geq '), ('∈', r'\in '),
    ('∞', r'\infty '), ('→', r'\rightarrow '), ('⇒', r'\Rightarrow '),
    ('≠', r'\neq '), ('≈', r'\approx '), ('°', r'^{\circ}'),
    # Arrows
    ('←', r'\leftarrow '), ('↑', r'\uparrow '), ('↓', r'\downarrow '),
    ('↔', r'\leftrightarrow '), ('⇐', r'\Leftarrow '), ('⇑', r'\Uparrow '),
]

# Unicode subscript characters
UNICODE_SUBSCRIPTS = {
    '₀': '_{0}', '₁': '_{1}', '₂': '_{2}', '₃': '_{3}', '₄': '_{4}',
    '₅': '_{5}', '₆': '_{6}', '₇': '_{7}', '₈': '_{8}', '₉': '_{9}',
    'ₐ': '_{a}', 'ₑ': '_{e}', 'ₕ': '_{h}', 'ᵢ': '_{i}',
    'ⱼ': '_{j}', 'ₖ': '_{k}', 'ₗ': '_{l}', 'ₘ': '_{m}',
    'ₙ': '_{n}', 'ₒ': '_{o}', 'ₚ': '_{p}', 'ᵣ': '_{r}',
    'ₛ': '_{s}', 'ₜ': '_{t}', 'ᵤ': '_{u}', 'ᵥ': '_{v}',
    'ₓ': '_{x}', '₊': '_{+}', '₋': '_{-}',
}

# Unicode superscript characters
UNICODE_SUPERSCRIPTS = {
    '⁰': '^{0}', '¹': '^{1}', '²': '^{2}', '³': '^{3}', '⁴': '^{4}',
    '⁵': '^{5}', '⁶': '^{6}', '⁷': '^{7}', '⁸': '^{8}', '⁹': '^{9}',
    '⁺': '^{+}', '⁻': '^{-}', '⁼': '^{=}',
    'ⁿ': '^{n}', 'ⁱ': '^{i}',
}

# Named subscripts that should use LaTeX operators or text
NAMED_SUBSCRIPTS = [
    (r'_max\b', r'_{\\max}'),
    (r'_min\b', r'_{\\min}'),
    (r'_total\b', r'_{\\mathrm{total}}'),
    (r'_ref\b', r'_{\\mathrm{ref}}'),
    (r'_th\b', r'_{\\mathrm{th}}'),
    (r'_sat\b', r'_{\\mathrm{sat}}'),
    (r'_avg\b', r'_{\\mathrm{avg}}'),
    (r'_in\b', r'_{\\mathrm{in}}'),
    (r'_out\b', r'_{\\mathrm{out}}'),
    (r'_crit\b', r'_{\\mathrm{crit}}'),
    (r'_opt\b', r'_{\\mathrm{opt}}'),
    (r'_init\b', r'_{\\mathrm{init}}'),
    (r'_jc\b', r'_{\\mathrm{jc}}'),
    (r'_hs\b', r'_{\\mathrm{hs}}'),
]

# Named operators / text phrases
NAMED_OPERATORS = [
    (' min:', r' \min '),
    (' max:', r' \max '),
    ('subject to:', r' \text{subject to:} '),
    ('subject to ', r' \text{subject to:} '),
    ('s.t.', r' \text{s.t.} '),
]


# ===========================================================================
# Unicode → LaTeX conversion
# ===========================================================================

def unicode_to_latex(text):
    """Convert Unicode math symbols and sub/superscripts to LaTeX commands.

    Processing order (critical — each step builds on the previous):
      1. Unicode sub/superscript chars → LaTeX {}-wrapped
      2. Named operators (min:, subject to:) → LaTeX commands
      3. Unicode math symbols → LaTeX commands
      4. Named subscripts (_max, _min...) → proper LaTeX
      5. General ASCII _ → _{...} wrapping
      6. General ASCII ^ → ^{...} wrapping
      7. Consecutive subscript joining (targeted, Bug #2 fix)

    Returns a LaTeX string suitable for latex2word.
    """
    s = text

    # Step 1: Unicode sub/superscript characters
    for sub, latex in UNICODE_SUBSCRIPTS.items():
        s = s.replace(sub, latex)
    for sup, latex in UNICODE_SUPERSCRIPTS.items():
        s = s.replace(sup, latex)

    # Step 2: Named operators (before symbol conversion)
    for pattern, replacement in NAMED_OPERATORS:
        s = s.replace(pattern, replacement)

    # Step 3: Unicode math symbols (longest match first)
    for uni, latex in sorted(UNICODE_MATH_SYMBOLS, key=lambda x: -len(x[0])):
        s = s.replace(uni, latex)

    # Step 4: Named subscripts (before general _ processing)
    for pattern, replacement in NAMED_SUBSCRIPTS:
        s = re.sub(pattern, replacement, s)

    # Step 5: General ASCII _ subscript wrapping
    s = re.sub(r'_([a-zA-Z]{2,})', r'_{\1}', s)
    s = re.sub(r'_(\\[a-zA-Z]+)', r'{\1}', s)
    s = re.sub(r'_([a-zA-Z])', r'_{\1}', s)
    s = re.sub(r'_(\d+)', r'_{\1}', s)

    # Step 6: General ASCII ^ superscript wrapping
    s = re.sub(r'\^([a-zA-Z]{2,})', r'^{\1}', s)
    s = re.sub(r'\^(\\[a-zA-Z]+)', r'^{\1}', s)
    s = re.sub(r'\^([a-zA-Z])', r'^{\1}', s)
    s = re.sub(r'\^(-?\d+)', r'^{\1}', s)

    # Step 7: Merge consecutive same-level subscripts (Bug #2 fix)
    # Only merges word_{X}_{Y} → word_{X,Y} where X,Y contain no nested braces.
    # Does NOT touch a_{b_{c}} (nested subscripts).
    s = re.sub(
        r'([a-zA-Z\\]+)_\{([^{}]+)\}_\{([^{}]+)\}',
        r'\1_{\2,\3}',
        s
    )

    # Final cleanup
    s = re.sub(r'\{(\{[^}]+\})\}', r'\1', s)  # remove doubled braces
    s = re.sub(r'\s+\{', '{', s)               # spaces before braces
    s = re.sub(r'\}\s+', '}', s)               # spaces after braces
    s = re.sub(r' +', ' ', s)                  # collapse spaces

    return s.strip()


# ===========================================================================
# Equation line detection
# ===========================================================================

def is_equation_line(line):
    """Detect if a stripped line is an equation (has math operators)
    or a text constraint that merely has a trailing equation number."""
    if not line:
        return False
    return bool(re.search(r'[=≤≥×·]|\\[a-zA-Z]|[_^]', line))


def parse_equation_line(line):
    """Parse a single line from a code block as a potential equation.

    Returns:
        ('eq', latex_str, eq_num) for equations
        ('text_eq', text, eq_num) for text constraints with equation numbers
        (None, None, None) if the line should be skipped
    """
    line = line.strip()
    if not line:
        return None, None, None

    # Extract trailing equation number e.g. "    (4)" or "    (8-10)"
    eq_num = None
    m = re.search(r'\s*\((\d+(?:[–\-]\d+)?)\)\s*$', line)
    if m:
        eq_num = m.group(1)
        line = line[:m.start()].strip()

    if not line:
        return 'text_eq', '', eq_num

    if is_equation_line(line):
        latex = unicode_to_latex(line)
        return 'eq', latex, eq_num
    else:
        return 'text_eq', line, eq_num


def parse_equation_block(code_lines):
    """Process a multi-line code block into individual equation items."""
    results = []
    for raw_line in code_lines:
        etype, content, eq_num = parse_equation_line(raw_line)
        if etype is not None:
            results.append((etype, content, eq_num))
    return results


# ===========================================================================
# OMML equation rendering
# ===========================================================================

def build_ommath_para(omml_element):
    """Wrap an m:oMath element in a properly structured m:oMathPara (display eq)."""
    omp = etree.Element(f"{{{MATH_NS}}}oMathPara")
    omp_pr = etree.SubElement(omp, f"{{{MATH_NS}}}oMathParaPr")
    jc = etree.SubElement(omp_pr, f"{{{MATH_NS}}}jc")
    jc.set(f"{{{MATH_NS}}}val", "center")
    om = etree.SubElement(omp, f"{{{MATH_NS}}}oMath")
    for child in omml_element:
        om.append(child)
    return omp


def append_omml(para, latex_str, display=True):
    """Convert LaTeX to OMML and append to paragraph.
    Returns the OMML element (oMathPara for display, oMath for inline)."""
    if not HAS_LATEX2WORD:
        run = para.add_run(f"[Equation: {latex_str}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return None
    try:
        eq = LatexToWordElement(latex_str)
        omml = eq.element()
        if display:
            omp = build_ommath_para(omml)
            para._element.append(omp)
            return omp
        else:
            para._element.append(omml)
            return omml
    except Exception:
        run = para.add_run(f"[Eq: {latex_str[:60]}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return None


def _compute_tab_stops(column_width_cm):
    """Compute center and right tab stops for equation numbering layout."""
    center_tab = column_width_cm / 2.0
    right_tab = column_width_cm
    return Cm(center_tab), Cm(right_tab)


def add_display_eq(para, latex_str, eq_cfg, eq_num, ctx, column_width_cm=None):
    """Add a display equation.
    Numbered: tab-stop layout with equation centered, number right-aligned.
    Unnumbered: paragraph centered, OMML only.
    """
    if column_width_cm is None:
        column_width_cm = 15.92  # A4 single-column fallback

    if eq_cfg.get("numbering"):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        center_stop, right_stop = _compute_tab_stops(column_width_cm)
        pf.tab_stops.add_tab_stop(center_stop, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        pf.tab_stops.add_tab_stop(right_stop, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        para.add_run("\t")
        append_omml(para, latex_str.strip(), display=True)
        fmt = eq_cfg.get("numbering_format", "({n})")
        num_text = fmt.replace("{n}", str(eq_num))
        run = para.add_run(f"\t{num_text}")
        nfont = eq_cfg.get("numbering_font", "Times New Roman")
        nsize = to_pt(eq_cfg.get("numbering_size", 12))
        set_run_font(run, nfont, nfont, nsize)
        bm_id = ctx.next_bookmark(f"eq{eq_num}")
        add_bookmark_to_run(run, f"eq{eq_num}", bm_id)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_omml(para, latex_str.strip(), display=True)


# ===========================================================================
# Bookmarks and hyperlinks
# ===========================================================================

def add_bookmark_to_run(run, name, bmid):
    """Wrap a run element with bookmark start/end tags."""
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bmid))
    bs.set(qn('w:name'), name)
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bmid))
    run._r.addprevious(bs)
    run._r.addnext(be)


def add_bookmark_to_para(para, name, bmid):
    """Add bookmark start/end to a paragraph element."""
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bmid))
    bs.set(qn('w:name'), name)
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bmid))
    para._p.append(bs)
    para._p.append(be)


def add_internal_hyperlink(para, anchor, display_text, font_cn, font_west, size):
    """Add a clickable internal cross-reference (e.g. '表 1')."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rs = OxmlElement('w:rStyle')
    rs.set(qn('w:val'), 'Hyperlink')
    rPr.append(rs)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(round(to_pt(size) * 2)))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = display_text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    hl.append(r)
    para._p.append(hl)
